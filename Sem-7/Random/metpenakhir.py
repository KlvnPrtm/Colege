from docx import Document

# Create a new Document for the resume
doc = Document()

# Title and Introduction
doc.add_heading("Resume Penelitian: Pengaruh Penggunaan Gadget Berlebihan terhadap Remaja", level=1)

# Cover Page
doc.add_page_break()
doc.add_heading("Resume Penelitian", level=1)
doc.add_paragraph("Disusun oleh:")
doc.add_paragraph("[Nama Anda]")
doc.add_paragraph("Mata Kuliah: Metode Penelitian")
doc.add_paragraph("Fakultas Teknologi Informasi")
doc.add_paragraph("[Universitas Anda]")
doc.add_paragraph("2024")
doc.add_page_break()

# Content: Resume
doc.add_heading("Resume", level=1)

# Introduction
doc.add_heading("1. Pendahuluan", level=2)
doc.add_paragraph("Penelitian ini dimulai dari pengamatan terhadap fenomena yang berkembang di kalangan remaja, "
                  "yaitu penggunaan gadget yang semakin meningkat seiring dengan kemajuan teknologi informasi. Gadget "
                  "seperti smartphone, tablet, dan laptop telah menjadi alat utama yang digunakan remaja untuk bersosialisasi, "
                  "menghibur diri, dan belajar. Namun, penggunaan yang berlebihan menimbulkan berbagai dampak negatif, "
                  "seperti terganggunya interaksi sosial, kesehatan mental, dan kualitas tidur mereka.")

# Topic Selection and Challenges
doc.add_heading("2. Penentuan Topik dan Tantangan", level=2)
doc.add_paragraph("Topik penelitian ini, yaitu 'Pengaruh Penggunaan Gadget Berlebihan terhadap Remaja,' dipilih karena relevansi dan urgensinya. "
                  "Beberapa tantangan yang muncul dalam proses ini adalah:")
doc.add_paragraph("- Menentukan batasan masalah yang spesifik untuk memastikan penelitian tetap fokus.")
doc.add_paragraph("- Mengembangkan instrumen penelitian, seperti kuesioner, yang sesuai dengan tujuan penelitian.")
doc.add_paragraph("Solusi untuk tantangan ini adalah melakukan kajian pustaka mendalam dan pilot study untuk menguji validitas instrumen.")

# Methodology
doc.add_heading("3. Metodologi Penelitian", level=2)
doc.add_paragraph("Penelitian ini menggunakan metode kuantitatif dengan pendekatan survei. Pengumpulan data dilakukan melalui kuesioner daring "
                  "dan wawancara mendalam. Responden yang menjadi target penelitian adalah remaja berusia 13–19 tahun, dengan jumlah sampel sebanyak 150 orang. "
                  "Data dianalisis menggunakan teknik statistik deskriptif, seperti rata-rata, median, dan standar deviasi.")

# Data Processing and Analysis
doc.add_heading("4. Pengolahan dan Analisis Data", level=2)
doc.add_paragraph("Setelah data terkumpul, dilakukan validasi untuk memastikan konsistensi dan kelengkapan data. Data kualitatif dikodekan menjadi kuantitatif "
                  "untuk memudahkan analisis. Hasil analisis menunjukkan bahwa:")
doc.add_paragraph("- Sebagian besar remaja menggunakan gadget selama 4-6 jam per hari, dengan waktu penggunaan dominan pada malam hari.")
doc.add_paragraph("- Penggunaan gadget berdampak pada menurunnya frekuensi interaksi sosial tatap muka.")
doc.add_paragraph("- Sebanyak 50 persen responden mengalami kesulitan tidur setelah menggunakan gadget lebih dari 2 jam sebelum tidur.")

# Highlights
doc.add_heading("5. Hal Menarik dan Menantang", level=2)
doc.add_paragraph("Beberapa hal menarik yang ditemukan selama penelitian ini adalah:")
doc.add_paragraph("- Kecenderungan remaja untuk menggunakan gadget sebagai sarana pelarian dari tekanan sosial di dunia nyata.")
doc.add_paragraph("- Dampak signifikan penggunaan gadget terhadap pola tidur, yang sering diabaikan oleh remaja.")
doc.add_paragraph("Tantangan utama adalah memastikan validitas kuesioner dan mendapatkan sampel yang representatif. Pilot study dan revisi kuesioner menjadi solusi efektif untuk mengatasi tantangan ini.")

# Conclusion and Hypothesis
doc.add_heading("6. Kesimpulan dan Hipotesis", level=2)
doc.add_paragraph("Berdasarkan hasil penelitian, dapat disimpulkan bahwa penggunaan gadget yang berlebihan memiliki dampak signifikan terhadap kehidupan remaja, terutama dalam aspek interaksi sosial, kesehatan mental, dan kualitas tidur. "
                  "Hipotesis yang dapat diajukan adalah bahwa pengurangan durasi penggunaan")

# Completing the resume content in Word format
# This portion continues the work interrupted previously

# Conclusion and Hypothesis (continuation)
doc.add_paragraph("Hipotesis yang dapat diajukan adalah bahwa pengurangan durasi penggunaan gadget, terutama pada malam hari, dapat secara signifikan meningkatkan kualitas tidur dan memperbaiki interaksi sosial remaja.")

# Reflections
doc.add_heading("7. Refleksi", level=2)
doc.add_paragraph("Penelitian ini memberikan pengalaman berharga dalam mengembangkan instrumen penelitian, mengumpulkan data, dan menganalisis hasil. "
                  "Hal ini juga menunjukkan pentingnya kolaborasi antara peneliti, responden, dan institusi terkait untuk menghasilkan penelitian yang berkualitas. "
                  "Meskipun banyak tantangan yang dihadapi, seperti keterbatasan waktu dan kompleksitas data, solusi inovatif seperti pilot study dan pengolahan data dengan alat statistik membantu menyelesaikan penelitian ini dengan baik.")

# Closing
doc.add_page_break()
doc.add_heading("Ucapan Terima Kasih", level=1)
doc.add_paragraph("Saya mengucapkan terima kasih kepada semua pihak yang telah membantu dalam proses penelitian ini, mulai dari dosen pembimbing, teman-teman, hingga para responden yang telah bersedia memberikan data mereka. Penelitian ini merupakan pengalaman berharga yang akan menjadi dasar untuk penelitian-penelitian selanjutnya.")

# Save the final document
final_resume_path = "Resume_Penelitian_Penggunaan_Gadget.docx"
doc.save(final_resume_path)
final_resume_path

# Attempting to re-save and confirm the path for download
final_resume_path_retry = "Resume_Penelitian_Penggunaan_Gadget_Final.docx"
doc.save(final_resume_path_retry)
final_resume_path_retry
